import ast
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger


class DocumentConfig:
    """
    Конфигурационный класс для управления путями сохранения файлов .docx.
    """

    def __init__(
        self,
        docx_output_path=None,
    ) -> None:
        """Инициализирует конфигурацию с заданным путём или использует путь по умолчанию.

        Args:
            docx_output_path (str | None): Пользовательский путь для сохранения .docx файлов.
        """
        self._docx_output_path = (
            docx_output_path or "/CONSPECTIUS/app/received_txt/{}.docx"
        )

    @property
    def docx_output_path(self):
        return self._docx_output_path

    @docx_output_path.setter
    def docx_output_path(self, new_docx_output_path: str):
        """Сеттер для изменения пути"""
        if not isinstance(new_docx_output_path, str):
            raise TypeError("Текст должен быть строкой.")
        self._docx_output_path = new_docx_output_path


class DocumentManager:
    """Менеджер по работе с файлами .docx"""

    def __init__(self, config: DocumentConfig) -> None:
        self._config = config
        self._path_docx = None

    @property
    def config(self):
        return self._config

    @property
    def path_docx(self):
        return self._path_docx

    @path_docx.setter
    def path_docx(self, file_path: str):
        self._path_docx = file_path

    def txt_to_docx(
        self,
        text: dataclass,
        lenght_conspect: str,
        telegram_id: int = None,
        new_file_title: int = None,
    ) -> None:
        """Сохраняет текст в файл по определённому пути и изменяет приватный атрибут _path_docx на значение пути нового файла

        Args:
            text (str): Dataclass, значение которого мы записываем в файл
        """
        title = str(telegram_id) if telegram_id is not None else new_file_title

        # Создаем объект документа
        doc = Document()
        heading = doc.add_heading(text.title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if lenght_conspect == "low":
            doc.add_paragraph(text.summary)
        else:
            doc.add_heading("Основные термины и понятия.", level=1)
            try:
                for term, value in ast.literal_eval(
                    text.key_terms_and_concepts
                ).items():
                    term = term[0].upper() + term[1:]
                    paragraph = doc.add_paragraph(style="List Number")
                    # Добавляем первую часть – жирный и курсивный текст
                    run1 = paragraph.add_run(term)
                    run1.bold = True
                    # Добавляем вторую часть – обычный текст
                    run2 = paragraph.add_run(f" - {value}")
            except Exception as err:
                doc.add_paragraph("Не получилось красиво оформить термины :(")
                logger.exception(
                    f"Не получилось красиво оформить термины: {err}"
                )
                doc.add_paragraph(text.key_terms_and_concepts)
                print(text.key_terms_and_concepts)

            doc.add_heading("Хронологический конспект лекции.", level=1)
            doc.add_paragraph(text.chronological_lecture_outline)

            doc.add_heading("Тест для понимая лекции", level=1)
            doc.add_paragraph(text.mini_test)

            ending = doc.add_paragraph(
                "Made with CONSPECTIUS. Made with love ❤️"
            )
            ending.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Полный путь к файлу
        file_path: str = self.config.docx_output_path.format(title)

        # Сохраняем документ в файл с расширением .docx
        doc.save(file_path)

        self.path_docx = file_path


# # Пример использования
# doc_config = DocumentConfig()
# doc_manager = DocumentManager(doc_config)
# doc_manager.txt_to_docx(
#     "Текст", 88417414
# ) # -> Новый файл с в диретории received_txt
# print(doc_manager.path_docx)
